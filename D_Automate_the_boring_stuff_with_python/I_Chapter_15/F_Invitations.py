

import docx
from docx.shared import Pt
from docx.enum.text import WD_BREAK

guest_names=open('FF_Guests.txt','r')
guest_list=[]

for line in guest_names:
    guest_list.append(line.strip('\n'))


invitations=docx.Document('FF_Custom_invitations.docx')
paragraph=0

for guest in guest_list:
    #adding first line
    invitations.paragraphs[paragraph].add_run('Nos gustaria contar con la placentera presencia de')
    invitations.paragraphs[paragraph].runs[0].add_break()

    #inserint guest from guest list
    invitations.paragraphs[paragraph].add_run(f'{guest}')
    guest_name=invitations.paragraphs[paragraph].runs[1]
    guest_name.font.name='Calibri'
    guest_name.font.size= Pt(35)
    guest_name.font.bold=True
    guest_name.font.imprint=True
    guest_name.add_break()

    #adding third line
    invitations.paragraphs[paragraph].add_run('en el Boulevard Santiago Mandrades la noche del')
    invitations.paragraphs[paragraph].runs[2].add_break()

    #adding forth line
    fecha=invitations.paragraphs[paragraph].add_run('24 de Dic')
    fecha.font.name='Calibri'
    fecha.font.size=Pt(25)
    fecha.add_break()
    #adding fifth line
    invitations.paragraphs[paragraph].add_run('a partir de las 9:00 PM').add_break()

    #Added this conditional to avoid a blank page at the end
    if (paragraph+1)<len(guest_list):
        invitations.paragraphs[paragraph].runs[4].add_break(WD_BREAK.PAGE)
        invitations.add_paragraph().style='Chavez24'
        paragraph+=1


    invitations.save('FF_Custom_invitations1.docx')
